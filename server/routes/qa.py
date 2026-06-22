"""智能问答模块路由 /api/qa（v3：流式 + 溯源 + 自动标题 + Word 导出）。"""
import io
import json

from flask import Blueprint, request, g, Response, send_file

from db import get_conn
from utils.response import ok, fail
from utils.jwt_util import token_required
from algo_context import llm, neo4j
from algo.graphrag.pipeline import graphrag_answer, retrieve_for_stream

bp = Blueprint("qa", __name__, url_prefix="/api/qa")


# ---------------- 会话 ----------------
@bp.post("/conversation")
@token_required
def create_conversation():
    data = request.get_json(force=True) or {}
    title = (data.get("title") or "新会话").strip()
    with get_conn() as conn:
        cur = conn.execute("INSERT INTO conversation(user_id,title) VALUES (?,?)",
                           (g.user_id, title))
        return ok({"id": cur.lastrowid, "title": title}, "会话已创建")


@bp.get("/conversation/list")
@token_required
def list_conversations():
    with get_conn() as conn:
        rows = conn.execute(
            "SELECT id,title,create_time FROM conversation WHERE user_id=? "
            "ORDER BY create_time DESC", (g.user_id,)).fetchall()
    return ok([{"id": r["id"], "title": r["title"], "createTime": r["create_time"]}
               for r in rows])


@bp.delete("/conversation/<int:cid>")
@token_required
def delete_conversation(cid):
    with get_conn() as conn:
        conn.execute("DELETE FROM qa_history WHERE conversation_id=?", (cid,))
        conn.execute("DELETE FROM conversation WHERE id=? AND user_id=?", (cid, g.user_id))
    return ok(msg="会话已删除")


@bp.put("/conversation/<int:cid>")
@token_required
def rename_conversation(cid):
    title = ((request.get_json(force=True) or {}).get("title") or "").strip()
    if not title:
        return fail("标题不能为空")
    with get_conn() as conn:
        conn.execute("UPDATE conversation SET title=? WHERE id=? AND user_id=?",
                     (title, cid, g.user_id))
    return ok(msg="已重命名")


@bp.get("/history/<int:cid>")
@token_required
def history(cid):
    with get_conn() as conn:
        rows = conn.execute(
            "SELECT question,answer,entities,create_time FROM qa_history "
            "WHERE conversation_id=? ORDER BY id ASC", (cid,)).fetchall()
    return ok([{"question": r["question"], "answer": r["answer"],
                "entities": r["entities"], "createTime": r["create_time"]} for r in rows])


@bp.post("/conversation/<int:cid>/title")
@token_required
def regen_title(cid):
    """生成标题：用最近一条提问让大模型重新生成会话标题。"""
    with get_conn() as conn:
        row = conn.execute(
            "SELECT question FROM qa_history WHERE conversation_id=? ORDER BY id DESC LIMIT 1",
            (cid,)).fetchone()
        if not row:
            return fail("会话暂无内容，无法生成标题")
        title = llm.gen_title(row["question"])
        conn.execute("UPDATE conversation SET title=? WHERE id=? AND user_id=?",
                     (title, cid, g.user_id))
    return ok({"title": title}, "标题已生成")


@bp.post("/conversation/<int:cid>/clear")
@token_required
def clear_records(cid):
    """清除记录：清空该会话的问答历史（保留会话本身）。"""
    with get_conn() as conn:
        conn.execute("DELETE FROM qa_history WHERE conversation_id=?", (cid,))
    return ok(msg="记录已清除")


@bp.post("/chat/image")
@token_required
def chat_image():
    """多模态问答：上传疾病图片 + 文本问题（qwen-vl）。"""
    question = (request.form.get("question") or "请识别这张医学相关图片并给出健康知识说明").strip()
    conversation_id = request.form.get("conversationId") or None
    f = request.files.get("image")
    if not f:
        return fail("未接收到图片")
    cid = _ensure_conversation(g.user_id, conversation_id, question)
    if not llm.available:
        answer = "图片识别需要配置通义千问多模态(qwen-vl)的 API-KEY；当前为演示模式，暂不能识别图片内容。"
        _save_qa(cid, g.user_id, "[图片] " + question, answer, [])
        return ok({"answer": answer, "conversationId": cid, "entities": [],
                   "citations": [], "reasoning": [], "tripleCount": 0,
                   "related": {"diseases": [], "records": []}})
    import base64
    img_b64 = base64.b64encode(f.read()).decode()
    mime = f.mimetype or "image/jpeg"
    # 先用问题做一次图谱检索，把上下文一并喂给多模态模型
    from algo.graphrag.pipeline import retrieve, build_context
    matched, triples, reasoning = retrieve(question, neo4j) if neo4j.ping() else ([], [], [])
    context = build_context(triples) if triples else ""
    try:
        answer = llm.answer_with_image(question, img_b64, mime, context)
    except Exception as e:
        answer = f"图片识别失败：{e}"
    entities = [m["name"] for m in matched][:8]
    _save_qa(cid, g.user_id, "[图片] " + question, answer, entities)
    return ok({"answer": answer, "conversationId": cid, "entities": entities,
               "citations": sorted({t.get("source") for t in triples if t.get("source")}),
               "reasoning": reasoning, "tripleCount": len(triples),
               "related": _related(entities)})


def _ensure_conversation(user_id, conversation_id, question):
    """无会话则自动创建，并用大模型生成标题。"""
    if conversation_id:
        return conversation_id
    title = llm.gen_title(question)
    with get_conn() as conn:
        cur = conn.execute("INSERT INTO conversation(user_id,title) VALUES (?,?)",
                           (user_id, title))
        return cur.lastrowid


def _save_qa(conversation_id, user_id, question, answer, entities):
    with get_conn() as conn:
        conn.execute(
            "INSERT INTO qa_history(conversation_id,user_id,question,answer,entities) "
            "VALUES (?,?,?,?,?)",
            (conversation_id, user_id, question, answer, "、".join(entities)))


def _related(entities):
    if not entities:
        return {"diseases": [], "records": []}
    with get_conn() as conn:
        diseases, records = [], []
        for name in entities:
            for r in conn.execute(
                    "SELECT id,product_code,name,cause,diagnosis_path FROM product_info "
                    "WHERE name LIKE ?", (f"%{name}%",)).fetchall():
                diseases.append(dict(r))
                for rec in conn.execute(
                        "SELECT batch_number,production_line,status FROM production_record "
                        "WHERE product_id=?", (r["id"],)).fetchall():
                    records.append(dict(rec))
    diseases = list({d["id"]: d for d in diseases}.values())
    return {"diseases": diseases, "records": records}


def _history_messages(cid, limit=4):
    """取该会话最近的问答对，构造多轮上下文（oldest→newest）。"""
    if not cid:
        return []
    with get_conn() as conn:
        rows = conn.execute(
            "SELECT question,answer FROM qa_history WHERE conversation_id=? "
            "ORDER BY id DESC LIMIT ?", (cid, limit)).fetchall()
    msgs = []
    for r in reversed(rows):
        msgs.append({"role": "user", "content": r["question"]})
        if r["answer"]:
            msgs.append({"role": "assistant", "content": r["answer"][:1500]})
    return msgs


def _last_entities(cid):
    """上一轮命中的实体（用于指代/追问时的实体复用）。"""
    if not cid:
        return []
    with get_conn() as conn:
        row = conn.execute(
            "SELECT entities FROM qa_history WHERE conversation_id=? ORDER BY id DESC LIMIT 1",
            (cid,)).fetchone()
    return [e for e in (row["entities"].split("、") if row and row["entities"] else []) if e]


# ---------------- 非流式问答 ----------------
@bp.post("/chat")
@token_required
def chat():
    data = request.get_json(force=True) or {}
    question = (data.get("question") or "").strip()
    if not question:
        return fail("问题不能为空")
    if not neo4j.ping():
        return fail("Neo4j 未连接，无法检索知识图谱", 500)
    prev_cid = data.get("conversationId")
    history = _history_messages(prev_cid)
    prev_entities = _last_entities(prev_cid)
    cid = _ensure_conversation(g.user_id, prev_cid, question)
    result = graphrag_answer(question, llm, neo4j, history=history, prev_entities=prev_entities)
    result["related"] = _related(result["entities"])
    result["conversationId"] = cid
    _save_qa(cid, g.user_id, question, result["answer"], result["entities"])
    return ok(result)


# ---------------- 流式问答 (SSE) ----------------
@bp.post("/chat/stream")
@token_required
def chat_stream():
    data = request.get_json(force=True) or {}
    question = (data.get("question") or "").strip()
    user_id = g.user_id
    conversation_id = data.get("conversationId")
    if not question:
        return fail("问题不能为空")
    if not neo4j.ping():
        return fail("Neo4j 未连接", 500)

    cid = _ensure_conversation(user_id, conversation_id, question)
    history = _history_messages(conversation_id)
    prev_entities = _last_entities(conversation_id)
    context, meta = retrieve_for_stream(question, llm, neo4j, prev_entities=prev_entities)
    meta["related"] = _related(meta["entities"])
    meta["conversationId"] = cid

    def gen():
        # 1) 先推送检索元数据（实体/三元组/推理链路/溯源）
        yield _sse("meta", meta)
        full = []
        if meta["triples"]:
            if llm.available:
                try:
                    for delta in llm.answer_stream(question, context, history=history):
                        full.append(delta)
                        yield _sse("delta", {"text": delta})
                except Exception as e:
                    from algo.graphrag.pipeline import _template
                    txt = _template(meta["entities"], meta["triples"]) + f"\n\n(大模型调用失败：{e})"
                    full = [txt]
                    yield _sse("delta", {"text": txt})
            else:
                from algo.graphrag.pipeline import _template
                txt = _template(meta["entities"], meta["triples"])
                full = [txt]
                yield _sse("delta", {"text": txt})
        else:
            if llm.available:
                try:
                    for delta in llm.chat_freeform_stream(question, history=history):
                        full.append(delta)
                        yield _sse("delta", {"text": delta})
                except Exception as e:
                    txt = f"（连接大模型失败：{e}）"
                    full = [txt]
                    yield _sse("delta", {"text": txt})
            else:
                txt = "未能在知识图谱中检索到相关知识点，请上传相关文档并重建图谱后再试。"
                full = [txt]
                yield _sse("delta", {"text": txt})

        answer = "".join(full)
        _save_qa(cid, user_id, question, answer, meta["entities"])
        yield _sse("done", {"conversationId": cid})

    return Response(gen(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


def _sse(event, data):
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


# ---------------- 导出 Word ----------------
@bp.get("/export/<int:cid>")
@token_required
def export_docx(cid):
    import docx
    with get_conn() as conn:
        conv = conn.execute("SELECT title FROM conversation WHERE id=?", (cid,)).fetchone()
        rows = conn.execute(
            "SELECT question,answer,create_time FROM qa_history "
            "WHERE conversation_id=? ORDER BY id ASC", (cid,)).fetchall()
    doc = docx.Document()
    doc.add_heading(conv["title"] if conv else "问答记录", level=0)
    doc.add_paragraph("基于 GraphRAG 的医疗健康知识诊断系统 · 问答咨询记录")
    for r in rows:
        doc.add_heading(f"问：{r['question']}", level=2)
        doc.add_paragraph(r["answer"] or "")
        p = doc.add_paragraph(r["create_time"] or "")
        p.italic = True
    doc.add_paragraph("\n本记录仅供健康科普参考，不能替代执业医师的诊断与治疗。")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name=f"问答记录_{cid}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
